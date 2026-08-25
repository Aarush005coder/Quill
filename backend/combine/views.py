import os
import uuid
import tempfile
import base64
import mimetypes
from io import BytesIO

from django.conf import settings
from django.http import FileResponse
from django.utils import timezone

from rest_framework import status, permissions
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.views import APIView

from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.pdfgen import canvas

from PIL import Image, ImageEnhance, ImageOps, ImageFilter
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import pdfplumber

from docx import Document as DocxDocument

import fitz

from .models import CombineOperation
from .serializers import CombineOperationSerializer

from openpyxl import load_workbook
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer,
    PageBreak,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ✅ IMPORT ADDED FOR REAL-TIME UNIFIED HISTORY
from history.utils import save_to_history

# ═════════════════════════════════════════════════════════════════
# HELPERS
# ═════════════════════════════════════════════════════════════════


def save_uploaded_temp(uploaded_file, suffix=""):
    safe_name = os.path.basename(uploaded_file.name)
    path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_{safe_name}{suffix}")
    with open(path, "wb") as destination:
        for chunk in uploaded_file.chunks():
            destination.write(chunk)
    return path


def user_output_dir(user):
    path = os.path.join(settings.MEDIA_ROOT, "combine", str(user.id))
    os.makedirs(path, exist_ok=True)
    return path


def relative_media_path(path):
    return os.path.relpath(path, settings.MEDIA_ROOT)


def cleanup_file(path):
    if path and os.path.exists(path):
        try:
            os.remove(path)
        except OSError:
            pass


def clamp_int(value, minimum, maximum, default):
    try:
        value = int(value)
    except (TypeError, ValueError):
        value = default
    return max(minimum, min(maximum, value))


def hex_to_rgb01(value, default=(1.0, 1.0, 1.0)):
    try:
        value = str(value).strip().lstrip("#")
        if len(value) != 6:
            return default
        return tuple(int(value[i:i + 2], 16) / 255 for i in (0, 2, 4))
    except (TypeError, ValueError):
        return default


# ═════════════════════════════════════════════════════════════════
# PDF MERGE
# ═════════════════════════════════════════════════════════════════


class PDFMergeView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        files = request.FILES.getlist("files")

        if len(files) < 2:
            return Response(
                {"success": False, "message": "At least 2 PDF files required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        operation = CombineOperation.objects.create(
            user=request.user,
            operation_type="pdf_merge",
            input_count=len(files),
        )

        input_paths = []
        try:
            merger = PdfWriter()
            for uploaded in files:
                if not uploaded.name.lower().endswith(".pdf"):
                    continue
                path = save_uploaded_temp(uploaded)
                input_paths.append(path)
                reader = PdfReader(path)
                for page in reader.pages:
                    merger.add_page(page)

            if len(merger.pages) == 0:
                raise ValueError("No valid PDF pages found.")

            output_path = os.path.join(user_output_dir(request.user), f"merged_{uuid.uuid4().hex}.pdf")
            with open(output_path, "wb") as out:
                merger.write(out)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = "merged_document.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="PDF Merge",
                description=f"Merged {len(files)} PDF files into {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"files": len(files), "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({
                "success": True,
                "message": "PDFs merged successfully.",
                "download_url": f"/api/combine/{operation.id}/download/",
                "operation_id": str(operation.id),
            })
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Merge failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            for path in input_paths:
                cleanup_file(path)


# ═════════════════════════════════════════════════════════════════
# IMAGE MERGE
# ═════════════════════════════════════════════════════════════════


class ImageMergeView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        files = request.FILES.getlist("files")
        layout = request.data.get("layout", "vertical")
        quality = clamp_int(request.data.get("quality", 90), 1, 100, 90)

        if len(files) < 2:
            return Response({"success": False, "message": "At least 2 images required."}, status=status.HTTP_400_BAD_REQUEST)

        operation = CombineOperation.objects.create(
            user=request.user,
            operation_type="image_merge",
            input_count=len(files),
            options={"layout": layout, "quality": quality},
        )

        images = []
        try:
            for uploaded in files:
                with Image.open(uploaded) as source:
                    images.append(source.convert("RGB"))

            if not images:
                raise ValueError("No valid images found.")

            if layout == "horizontal":
                total_width = sum(img.width for img in images)
                max_height = max(img.height for img in images)
                merged = Image.new("RGB", (total_width, max_height), (255, 255, 255))
                x = 0
                for img in images:
                    merged.paste(img, (x, (max_height - img.height) // 2))
                    x += img.width
            elif layout == "grid":
                cols = clamp_int(request.data.get("cols", 2), 1, 10, 2)
                rows = (len(images) + cols - 1) // cols
                max_w = max(img.width for img in images)
                max_h = max(img.height for img in images)
                merged = Image.new("RGB", (max_w * cols, max_h * rows), (255, 255, 255))
                for index, img in enumerate(images):
                    x = (index % cols) * max_w + (max_w - img.width) // 2
                    y = (index // cols) * max_h + (max_h - img.height) // 2
                    merged.paste(img, (x, y))
            else:
                max_width = max(img.width for img in images)
                total_height = sum(img.height for img in images)
                merged = Image.new("RGB", (max_width, total_height), (255, 255, 255))
                y = 0
                for img in images:
                    merged.paste(img, ((max_width - img.width) // 2, y))
                    y += img.height

            output_path = os.path.join(user_output_dir(request.user), f"merged_images_{uuid.uuid4().hex}.jpg")
            merged.save(output_path, "JPEG", quality=quality, optimize=True)
            merged.close()

            operation.output_file = relative_media_path(output_path)
            operation.output_name = "merged_images.jpg"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Image Merge",
                description=f"Merged {len(files)} images into {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"files": len(files), "layout": layout, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({
                "success": True,
                "message": "Images merged successfully.",
                "download_url": f"/api/combine/{operation.id}/download/",
                "operation_id": str(operation.id),
            })
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Merge failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            for image in images:
                try:
                    image.close()
                except Exception:
                    pass


# ═════════════════════════════════════════════════════════════════
# IMAGE TO PDF
# ═════════════════════════════════════════════════════════════════


class ImageToPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        files = request.FILES.getlist("files")
        page_size = request.data.get("page_size", "A4")

        if not files:
            return Response({"success": False, "message": "At least 1 image required."}, status=status.HTTP_400_BAD_REQUEST)

        operation = CombineOperation.objects.create(user=request.user, operation_type="image_to_pdf", input_count=len(files))
        temp_images = []

        try:
            pagesize = {"A4": A4, "letter": letter}.get(page_size, A4)
            width, height = pagesize
            output_path = os.path.join(user_output_dir(request.user), f"images_to_pdf_{uuid.uuid4().hex}.pdf")
            pdf = canvas.Canvas(output_path, pagesize=pagesize)

            for uploaded in files:
                with Image.open(uploaded) as source:
                    img = source.convert("RGB")
                    img.thumbnail((width - 40, height - 40))
                    temp_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}.png")
                    img.save(temp_path, "PNG")
                    temp_images.append(temp_path)
                    pdf.drawImage(temp_path, (width - img.width) / 2, (height - img.height) / 2, width=img.width, height=img.height)
                    pdf.showPage()

            pdf.save()
            operation.output_file = relative_media_path(output_path)
            operation.output_name = "images_converted.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Image to PDF",
                description=f"Converted {len(files)} images to {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"files": len(files), "page_size": page_size, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({
                "success": True,
                "message": "Images converted to PDF.",
                "download_url": f"/api/combine/{operation.id}/download/",
                "operation_id": str(operation.id),
            })
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Conversion failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            for path in temp_images:
                cleanup_file(path)


# ═════════════════════════════════════════════════════════════════
# IMAGE CONVERTER
# ═════════════════════════════════════════════════════════════════


class ImageConvertView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    ALLOWED_FORMATS = {"jpg", "jpeg", "png", "webp", "gif", "bmp", "tiff", "ico", "svg"}
    FORMAT_EXTENSIONS = {"jpg": "jpg", "jpeg": "jpeg", "png": "png", "webp": "webp", "gif": "gif", "bmp": "bmp", "tiff": "tiff", "ico": "ico", "svg": "svg"}
    FORMAT_MIMES = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "webp": "image/webp", "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff", "ico": "image/x-icon", "svg": "image/svg+xml"}

    @staticmethod
    def prepare_image(image, target_format):
        target_format = target_format.lower()
        if target_format in {"jpg", "jpeg"}:
            if image.mode in {"RGBA", "LA", "P"}:
                rgba = image.convert("RGBA")
                background = Image.new("RGB", rgba.size, (255, 255, 255))
                background.paste(rgba, mask=rgba.getchannel("A"))
                return background
            return image.convert("RGB")
        if target_format == "png":
            if image.mode in {"RGBA", "LA"}: return image
            if image.mode == "P": return image.convert("RGBA")
            return image.convert("RGB")
        if target_format == "webp":
            if image.mode in {"RGBA", "LA"}: return image.convert("RGBA")
            return image.convert("RGB")
        if target_format == "gif":
            if image.mode == "P": return image
            return image.convert("P")
        if target_format == "bmp": return image.convert("RGB")
        if target_format == "tiff":
            if image.mode in {"RGBA", "RGB", "L"}: return image
            return image.convert("RGBA")
        if target_format == "ico": return image.convert("RGBA")
        return image.convert("RGB")

    @staticmethod
    def image_to_svg(image, original_bytes, original_mime):
        width, height = image.size
        png_buffer = BytesIO()
        svg_image = image if image.mode in {"RGB", "RGBA", "L"} else image.convert("RGBA")
        svg_image.save(png_buffer, format="PNG")
        encoded = base64.b64encode(png_buffer.getvalue()).decode("ascii")
        return f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" xmlns:xlink="http://www.w3.org/1999/xlink" width="{width}" height="{height}" viewBox="0 0 {width} {height}">
    <image width="{width}" height="{height}" preserveAspectRatio="none" href="data:image/png;base64,{encoded}" />
</svg>'''.encode("utf-8")

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "Image file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        target_format = str(request.data.get("format") or request.data.get("output_format") or "").strip().lower().replace(".", "")
        
        if target_format not in self.ALLOWED_FORMATS:
            return Response({"success": False, "message": "Unsupported output format."}, status=status.HTTP_400_BAD_REQUEST)

        quality = clamp_int(request.data.get("quality", 92), 1, 100, 92)
        operation = CombineOperation.objects.create(user=request.user, operation_type="image_convert", input_count=1, options={"format": target_format, "quality": quality})

        try:
            uploaded_bytes = uploaded.read()
            if not uploaded_bytes: raise ValueError("Uploaded image is empty.")

            if uploaded.name.lower().endswith(".svg"):
                if target_format == "svg":
                    output_bytes = uploaded_bytes
                else:
                    try:
                        import cairosvg
                    except ImportError:
                        raise RuntimeError("SVG input conversion requires 'cairosvg'. Install with: pip install cairosvg")
                    png_bytes = cairosvg.svg2png(bytestring=uploaded_bytes)
                    image = Image.open(BytesIO(png_bytes))
            else:
                image = Image.open(BytesIO(uploaded_bytes))

            image.load()
            output_buffer = BytesIO()

            if target_format == "svg":
                output_buffer.write(self.image_to_svg(image, uploaded_bytes, uploaded.content_type or "image/png"))
            else:
                processed = self.prepare_image(image, target_format)
                save_format_map = {"jpg": "JPEG", "jpeg": "JPEG", "png": "PNG", "webp": "WEBP", "gif": "GIF", "bmp": "BMP", "tiff": "TIFF", "ico": "ICO"}
                save_kwargs = {"quality": quality, "optimize": True, "progressive": True} if save_format_map[target_format] == "JPEG" else {"quality": quality, "method": 6} if save_format_map[target_format] == "WEBP" else {"optimize": True} if save_format_map[target_format] == "PNG" else {"sizes": [(16, 16), (32, 32), (48, 48), (64, 64), (128, 128), (256, 256)]} if save_format_map[target_format] == "ICO" else {}
                
                processed.save(output_buffer, format=save_format_map[target_format], **save_kwargs)
                if processed is not image: processed.close()

            image.close()
            output_bytes = output_buffer.getvalue()
            output_buffer.close()
            if not output_bytes: raise ValueError("Failed to generate converted image.")

            original_base = os.path.splitext(uploaded.name)[0]
            output_filename = f"{original_base}.{self.FORMAT_EXTENSIONS[target_format]}"
            output_path = os.path.join(user_output_dir(request.user), f"converted_{uuid.uuid4().hex}.{self.FORMAT_EXTENSIONS[target_format]}")
            
            with open(output_path, "wb") as out:
                out.write(output_bytes)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = output_filename
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Image Converter",
                description=f"Converted image to {target_format.upper()}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"toolType": "Image Converter", "format": target_format, "quality": quality, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({
                "success": True, "message": f"Image converted to {target_format.upper()} successfully.",
                "format": target_format, "mime_type": self.FORMAT_MIMES[target_format],
                "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id),
                "output_name": operation.output_name, "output_size": operation.output_size,
            }, status=status.HTTP_200_OK)

        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": "Image conversion failed.", "error": str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ═════════════════════════════════════════════════════════════════
# WORD MERGE
# ═════════════════════════════════════════════════════════════════


class WordMergeView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        files = request.FILES.getlist("files")
        if len(files) < 2:
            return Response({"success": False, "message": "At least 2 Word files required."}, status=status.HTTP_400_BAD_REQUEST)

        operation = CombineOperation.objects.create(user=request.user, operation_type="word_merge", input_count=len(files))
        temp_paths = []

        try:
            merged_doc = DocxDocument()
            for uploaded in files:
                if not uploaded.name.lower().endswith((".docx", ".doc")): continue
                path = save_uploaded_temp(uploaded)
                temp_paths.append(path)
                doc = DocxDocument(path)
                for element in doc.element.body:
                    merged_doc.element.body.append(element)

            output_path = os.path.join(user_output_dir(request.user), f"merged_word_{uuid.uuid4().hex}.docx")
            merged_doc.save(output_path)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = "merged_document.docx"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Word Merge",
                description=f"Merged {len(files)} Word documents into {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"files": len(files), "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "Word documents merged.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Merge failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            for path in temp_paths:
                cleanup_file(path)


# ═════════════════════════════════════════════════════════════════
# PDF TO WORD
# ═════════════════════════════════════════════════════════════════


class PDFToWordView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        operation = CombineOperation.objects.create(user=request.user, operation_type="pdf_to_word", input_count=1)
        temp_pdf = None

        try:
            temp_pdf = save_uploaded_temp(uploaded)
            reader = PdfReader(temp_pdf)
            doc = DocxDocument()

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    for paragraph in text.split("\n"):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph)
                    doc.add_paragraph()

            output_path = os.path.join(user_output_dir(request.user), f"converted_{uuid.uuid4().hex}.docx")
            doc.save(output_path)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + ".docx"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="PDF to Word",
                description=f"Converted {uploaded.name} to {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "PDF to Word", "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "PDF converted to Word.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Conversion failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_pdf)


# ═════════════════════════════════════════════════════════════════
# PDF TO EXCEL
# ═════════════════════════════════════════════════════════════════

class PDFToExcelView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]

        if not uploaded.name.lower().endswith(".pdf"):
            return Response({"success": False, "message": "Only PDF files are supported."}, status=status.HTTP_400_BAD_REQUEST)

        operation = CombineOperation.objects.create(
            user=request.user,
            operation_type="pdf_to_excel",
            input_count=1,
            options={"source_file": uploaded.name},
        )

        temp_pdf = None
        output_path = None

        try:
            temp_pdf = save_uploaded_temp(uploaded)
            workbook = Workbook()
            default_sheet = workbook.active
            workbook.remove(default_sheet)

            total_tables = 0
            total_rows = 0

            with pdfplumber.open(temp_pdf) as pdf:
                if not pdf.pages:
                    raise ValueError("The PDF contains no pages.")

                for page_index, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()

                    if tables:
                        for table_index, table in enumerate(tables, start=1):
                            if not table: continue
                            sheet_name = f"Page {page_index}" if len(tables) == 1 else f"Page {page_index} Table {table_index}"
                            sheet_name = sheet_name[:31]
                            worksheet = workbook.create_sheet(title=sheet_name)
                            total_tables += 1

                            for row_index, row in enumerate(table, start=1):
                                if row is None: continue
                                cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                if not any(str(value).strip() for value in cleaned_row): continue
                                total_rows += 1

                                for column_index, value in enumerate(cleaned_row, start=1):
                                    cell = worksheet.cell(row=row_index, column=column_index, value=value)
                                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                                    if row_index == 1:
                                        cell.font = Font(bold=True, color="FFFFFF")
                                        cell.fill = PatternFill(fill_type="solid", fgColor="2563EB")
                                
                                thin = Side(style="thin", color="D1D5DB")
                                for column_index in range(1, len(cleaned_row) + 1):
                                    worksheet.cell(row=row_index, column=column_index).border = Border(left=thin, right=thin, top=thin, bottom=thin)

                            for column_cells in worksheet.columns:
                                max_length = 0
                                column_letter = column_cells[0].column_letter
                                for cell in column_cells:
                                    try:
                                        length = len(str(cell.value) if cell.value is not None else "")
                                        if length > max_length: max_length = length
                                    except Exception: pass
                                worksheet.column_dimensions[column_letter].width = min(max(max_length + 2, 12), 50)
                            worksheet.freeze_panes = "A2"
                    else:
                        text = page.extract_text() or ""
                        if text.strip():
                            worksheet = workbook.create_sheet(title=f"Page {page_index}"[:31])
                            worksheet["A1"] = "Extracted Text"
                            worksheet["A1"].font = Font(bold=True, color="FFFFFF")
                            worksheet["A1"].fill = PatternFill(fill_type="solid", fgColor="2563EB")
                            worksheet["A1"].alignment = Alignment(vertical="top")
                            row_number = 2
                            for line in text.splitlines():
                                line = line.strip()
                                if not line: continue
                                worksheet.cell(row=row_number, column=1, value=line).alignment = Alignment(vertical="top", wrap_text=True)
                                row_number += 1
                                total_rows += 1
                            worksheet.column_dimensions["A"].width = 100
                            worksheet.freeze_panes = "A2"

            if len(workbook.worksheets) == 0:
                worksheet = workbook.create_sheet(title="PDF Content")
                worksheet["A1"] = "No extractable table or text was found in this PDF."
                worksheet.column_dimensions["A"].width = 80

            output_path = os.path.join(user_output_dir(request.user), f"pdf_to_excel_{uuid.uuid4().hex}.xlsx")
            workbook.save(output_path)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + ".xlsx"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.options = {"source_file": uploaded.name, "tables_extracted": total_tables, "rows_extracted": total_rows}
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="PDF to Excel",
                description=f"Converted {uploaded.name} to {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "PDF to Excel", "tables_extracted": total_tables, "rows_extracted": total_rows, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({
                "success": True, "message": "PDF converted to Excel successfully.",
                "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id),
                "output_name": operation.output_name, "output_size": operation.output_size,
                "tables_extracted": total_tables, "rows_extracted": total_rows,
            }, status=status.HTTP_200_OK)

        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": "PDF to Excel conversion failed.", "error": str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_pdf)


# ═════════════════════════════════════════════════════════════════
# WORD TO PDF
# ═════════════════════════════════════════════════════════════════


class WordToPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "Word file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        operation = CombineOperation.objects.create(user=request.user, operation_type="word_to_pdf", input_count=1)
        temp_docx = None

        try:
            temp_docx = save_uploaded_temp(uploaded)
            doc = DocxDocument(temp_docx)
            output_path = os.path.join(user_output_dir(request.user), f"converted_{uuid.uuid4().hex}.pdf")
            pdf = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            y = height - 50

            for paragraph in doc.paragraphs:
                if y < 50:
                    pdf.showPage()
                    y = height - 50
                pdf.drawString(50, y, paragraph.text[:100])
                y -= 15

            pdf.save()
            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + ".pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Word to PDF",
                description=f"Converted {uploaded.name} to {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "Word to PDF", "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "Word converted to PDF.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Conversion failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_docx)


# ═════════════════════════════════════════════════════════════════
# EXCEL TO PDF
# ═════════════════════════════════════════════════════════════════

class ExcelToPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    ALLOWED_EXTENSIONS = {".xlsx", ".xls"}

    @staticmethod
    def safe_text(value):
        if value is None: return ""
        if hasattr(value, "strftime"):
            try: return value.strftime("%Y-%m-%d %H:%M:%S")
            except Exception: return str(value)
        return str(value)

    @staticmethod
    def build_pdf(workbook, output_path, original_filename):
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("ExcelTitle", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, alignment=TA_CENTER, textColor=colors.HexColor("#1e293b"), spaceAfter=8)
        sheet_style = ParagraphStyle("SheetTitle", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, alignment=TA_LEFT, textColor=colors.HexColor("#334155"), spaceAfter=8)
        cell_style = ParagraphStyle("ExcelCell", parent=styles["BodyText"], fontName="Helvetica", fontSize=7, leading=9, alignment=TA_LEFT, textColor=colors.HexColor("#0f172a"))
        header_cell_style = ParagraphStyle("ExcelHeader", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=7, leading=9, alignment=TA_LEFT, textColor=colors.white)

        document = SimpleDocTemplate(output_path, pagesize=landscape(A4), rightMargin=10 * mm, leftMargin=10 * mm, topMargin=10 * mm, bottomMargin=10 * mm, title=f"Excel to PDF - {original_filename}", author="quill")
        page_width, page_height = landscape(A4)
        usable_width = page_width - 20 * mm
        story = []

        story.append(Paragraph("Excel Spreadsheet", title_style))
        story.append(Paragraph(original_filename, ParagraphStyle("FileName", parent=styles["BodyText"], fontSize=8, textColor=colors.HexColor("#64748b"), alignment=TA_CENTER, spaceAfter=14)))

        visible_sheets = [sheet for sheet in workbook.worksheets if sheet.max_row > 0 and sheet.max_column > 0]
        if not visible_sheets:
            story.append(Paragraph("The Excel file contains no printable data.", styles["BodyText"]))

        for sheet_index, ws in enumerate(visible_sheets):
            story.append(Paragraph(f"Sheet: {ws.title}", sheet_style))
            max_row, max_col = ws.max_row, ws.max_column

            while max_row > 1:
                if any(ws.cell(max_row, col).value is not None for col in range(1, max_col + 1)): break
                max_row -= 1
            while max_col > 1:
                if any(ws.cell(row, max_col).value is not None for row in range(1, max_row + 1)): break
                max_col -= 1

            column_widths = []
            for col_idx in range(1, max_col + 1):
                max_length = 5
                for row_idx in range(1, min(max_row, 100) + 1):
                    value = ws.cell(row_idx, col_idx).value
                    if value is not None: max_length = max(max_length, len(ExcelToPDFView.safe_text(value)))
                column_widths.append(min(max_length * 4.5, 30 * 4.5))

            total_width = sum(column_widths)
            if total_width > usable_width:
                scale = usable_width / total_width
                column_widths = [max(width * scale, 18) for width in column_widths]
            
            table_data = []
            for row_idx in range(1, max_row + 1):
                row_data = []
                for col_idx in range(1, max_col + 1):
                    value = ws.cell(row_idx, col_idx).value
                    text = ExcelToPDFView.safe_text(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    row_data.append(Paragraph(text or " ", header_cell_style if row_idx == 1 else cell_style))
                table_data.append(row_data)

            if table_data:
                table = Table(table_data, colWidths=column_widths, repeatRows=1, splitByRow=1, hAlign="LEFT")
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(table)
                if sheet_index < len(visible_sheets) - 1: story.append(PageBreak())

        document.build(story)

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "Excel file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        extension = os.path.splitext(uploaded.name)[1].lower()

        if extension not in self.ALLOWED_EXTENSIONS:
            return Response({"success": False, "message": "Only .xlsx and .xls files are supported."}, status=status.HTTP_400_BAD_REQUEST)

        operation = CombineOperation.objects.create(user=request.user, operation_type="excel_to_pdf", input_count=1, options={"source_filename": uploaded.name, "extension": extension})
        temp_excel = None

        try:
            temp_excel = save_uploaded_temp(uploaded)
            if extension == ".xlsx":
                workbook = load_workbook(temp_excel, data_only=True, read_only=False)
            else:
                import xlrd, pandas as pd
                excel_file = pd.ExcelFile(temp_excel, engine="xlrd")
                output_xlsx_temp = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_converted.xlsx")
                with pd.ExcelWriter(output_xlsx_temp, engine="openpyxl") as writer:
                    for sheet_name in excel_file.sheet_names:
                        pd.read_excel(excel_file, sheet_name=sheet_name, header=None).to_excel(writer, sheet_name=sheet_name[:31], index=False, header=False)
                workbook = load_workbook(output_xlsx_temp, data_only=True)
                cleanup_file(output_xlsx_temp)

            base_name = os.path.splitext(uploaded.name)[0]
            output_filename = f"{base_name}.pdf"
            output_path = os.path.join(user_output_dir(request.user), f"excel_to_pdf_{uuid.uuid4().hex}.pdf")

            self.build_pdf(workbook, output_path, uploaded.name)
            workbook.close()

            operation.output_file = relative_media_path(output_path)
            operation.output_name = output_filename
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Excel to PDF",
                description=f"Converted {uploaded.name} to {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "Excel to PDF", "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "Excel converted to PDF successfully.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id), "output_name": output_filename, "output_size": operation.output_size}, status=status.HTTP_200_OK)

        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": "Excel to PDF conversion failed.", "error": str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_excel)


# ═════════════════════════════════════════════════════════════════
# COMPRESS PDF
# ═════════════════════════════════════════════════════════════════


class CompressPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        quality = request.data.get("quality", "medium")
        operation = CombineOperation.objects.create(user=request.user, operation_type="compress_pdf", input_count=1, options={"quality": quality})
        temp_pdf = None

        try:
            temp_pdf = save_uploaded_temp(uploaded)
            reader = PdfReader(temp_pdf)
            writer = PdfWriter()

            for page in reader.pages:
                writer.add_page(page)
            writer.add_metadata({"/Producer": "quill"})

            output_path = os.path.join(user_output_dir(request.user), f"compressed_{uuid.uuid4().hex}.pdf")
            with open(output_path, "wb") as out:
                writer.write(out)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + "_compressed.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="PDF Compress",
                description=f"Compressed {uploaded.name} to {operation.output_name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "PDF Compress", "quality": quality, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "PDF compressed.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Compression failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_pdf)


# ═════════════════════════════════════════════════════════════════
# ROTATE PDF
# ═════════════════════════════════════════════════════════════════


class RotatePDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        rotation = clamp_int(request.data.get("rotation", 90), 0, 360, 90)
        pages = request.data.get("pages", "all")
        operation = CombineOperation.objects.create(user=request.user, operation_type="rotate_pdf", input_count=1, options={"rotation": rotation, "pages": pages})
        temp_pdf = None

        try:
            temp_pdf = save_uploaded_temp(uploaded)
            reader = PdfReader(temp_pdf)
            writer = PdfWriter()

            page_list = None if pages == "all" else [int(p.strip()) - 1 for p in str(pages).split(",") if p.strip()]

            for index, page in enumerate(reader.pages):
                if page_list is None or index in page_list:
                    page.rotate(rotation)
                writer.add_page(page)

            output_path = os.path.join(user_output_dir(request.user), f"rotated_{uuid.uuid4().hex}.pdf")
            with open(output_path, "wb") as out:
                writer.write(out)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + "_rotated.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Rotate PDF",
                description=f"Rotated {uploaded.name} by {rotation} degrees",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "Rotate PDF", "rotation": rotation, "pages": pages, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "PDF rotated.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Rotation failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_pdf)


# ═════════════════════════════════════════════════════════════════
# ORGANIZE PDF
# ═════════════════════════════════════════════════════════════════

class OrganizePDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        if not uploaded.name.lower().endswith(".pdf"):
            return Response({"success": False, "message": "Only PDF files are supported."}, status=status.HTTP_400_BAD_REQUEST)

        page_order_raw = request.data.get("page_order", "")
        if not str(page_order_raw).strip():
            return Response({"success": False, "message": "Page order is required."}, status=status.HTTP_400_BAD_REQUEST)

        operation = CombineOperation.objects.create(user=request.user, operation_type="organize_pdf", input_count=1, options={"page_order": str(page_order_raw)})
        temp_pdf = None

        try:
            temp_pdf = save_uploaded_temp(uploaded)
            reader = PdfReader(temp_pdf)
            writer = PdfWriter()
            total_pages = len(reader.pages)

            requested_pages = []
            for value in str(page_order_raw).split(","):
                value = value.strip()
                if not value: continue
                try:
                    page_number = int(value)
                    if 1 <= page_number <= total_pages: requested_pages.append(page_number)
                except ValueError: continue

            if not requested_pages: raise ValueError("No valid page numbers were provided.")

            seen = set()
            final_order = []
            for page_number in requested_pages:
                if page_number not in seen:
                    seen.add(page_number)
                    final_order.append(page_number)

            for page_number in final_order:
                writer.add_page(reader.pages[page_number - 1])

            if len(writer.pages) == 0: raise ValueError("No pages were selected.")

            output_path = os.path.join(user_output_dir(request.user), f"organized_{uuid.uuid4().hex}.pdf")
            with open(output_path, "wb") as output_file:
                writer.write(output_file)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + "_organized.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Organize PDF",
                description=f"Reorganized pages of {uploaded.name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "Organize PDF", "page_order": final_order, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "PDF pages organized successfully.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id), "output_name": operation.output_name, "output_size": operation.output_size, "page_order": final_order}, status=status.HTTP_200_OK)

        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Organize PDF failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_pdf)


# ═════════════════════════════════════════════════════════════════
# WATERMARK PDF
# ═════════════════════════════════════════════════════════════════


class WatermarkPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        watermark_text = request.data.get("watermark_text", "quill")
        watermark_size = request.data.get("watermark_size", "medium")
        watermark_pages = request.data.get("watermark_pages", "all")
        watermark_custom_pages = request.data.get("watermark_custom_pages", "")

        custom_page_numbers = set()
        if watermark_pages == "custom" and watermark_custom_pages:
            for part in str(watermark_custom_pages).split(","):
                part = part.strip()
                if "-" in part:
                    try:
                        start, end = map(int, part.split("-"))
                        custom_page_numbers.update(range(start, end + 1))
                    except ValueError: pass
                else:
                    try: custom_page_numbers.add(int(part))
                    except ValueError: pass

        operation = CombineOperation.objects.create(
            user=request.user, operation_type="watermark_pdf", input_count=1,
            options={"watermark_text": watermark_text, "watermark_size": watermark_size, "watermark_pages": watermark_pages, "watermark_custom_pages": watermark_custom_pages},
        )

        temp_pdf = None
        wm_templates = {}

        try:
            temp_pdf = save_uploaded_temp(uploaded)
            reader = PdfReader(temp_pdf)
            writer = PdfWriter()

            for size_type in ["small", "medium", "large"]:
                wm_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_wm_{size_type}.pdf")
                c = canvas.Canvas(wm_path, pagesize=letter)
                width, height = letter

                if size_type == "large":
                    c.setFont("Helvetica-Bold", 80)
                    c.setFillColor(colors.Color(0.7, 0.7, 0.7, alpha=0.3))
                    c.saveState()
                    c.translate(width / 2, height / 2)
                    c.rotate(45)
                    c.drawCentredString(0, 0, watermark_text)
                    c.restoreState()
                elif size_type == "medium":
                    c.setFont("Helvetica", 40)
                    c.setFillColor(colors.Color(0.7, 0.7, 0.7, alpha=0.25))
                    for x in range(0, int(width) + 200, 200):
                        for y in range(0, int(height) + 250, 250):
                            c.saveState()
                            c.translate(x, y)
                            c.rotate(45)
                            c.drawCentredString(0, 0, watermark_text)
                            c.restoreState()
                else:
                    c.setFont("Helvetica", 20)
                    c.setFillColor(colors.Color(0.7, 0.7, 0.7, alpha=0.15))
                    for x in range(0, int(width) + 100, 100):
                        for y in range(0, int(height) + 120, 120):
                            c.saveState()
                            c.translate(x, y)
                            c.rotate(45)
                            c.drawCentredString(0, 0, watermark_text)
                            c.restoreState()
                
                c.save()
                wm_templates[size_type] = wm_path

            wm_readers = {size_type: PdfReader(path) for size_type, path in wm_templates.items()}

            for page_index, page in enumerate(reader.pages):
                page_num = page_index + 1
                should_add_watermark = False

                if watermark_pages == "all": should_add_watermark = True
                elif watermark_pages == "even" and page_num % 2 == 0: should_add_watermark = True
                elif watermark_pages == "odd" and page_num % 2 != 0: should_add_watermark = True
                elif watermark_pages == "custom" and page_num in custom_page_numbers: should_add_watermark = True

                if should_add_watermark:
                    wm_page = wm_readers[watermark_size].pages[0]
                    page.merge_page(wm_page)
                writer.add_page(page)

            output_path = os.path.join(user_output_dir(request.user), f"watermarked_{uuid.uuid4().hex}.pdf")
            with open(output_path, "wb") as out:
                writer.write(out)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + "_watermarked.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Watermark PDF",
                description=f"Added watermark to {uploaded.name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "Watermark PDF", "watermark_text": watermark_text, "watermark_size": watermark_size, "watermark_pages": watermark_pages, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "Watermark added successfully.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})

        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Watermark failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_pdf)
            for path in wm_templates.values(): cleanup_file(path)


# ═════════════════════════════════════════════════════════════════
# SPLIT PDF
# ═════════════════════════════════════════════════════════════════


class SplitPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        split_mode = request.data.get("split_mode", "keep")
        split_pages_str = request.data.get("split_pages", "")
        split_parity = request.data.get("split_parity", "all")

        operation = CombineOperation.objects.create(
            user=request.user, operation_type="split_pdf", input_count=1,
            options={"split_mode": split_mode, "split_pages": split_pages_str, "split_parity": split_parity},
        )

        temp_pdf = None
        try:
            temp_pdf = save_uploaded_temp(uploaded)
            reader = PdfReader(temp_pdf)
            writer = PdfWriter()
            total_pages = len(reader.pages)

            target_pages = set()
            if split_pages_str.strip():
                for part in split_pages_str.split(','):
                    part = part.strip()
                    if '-' in part:
                        try:
                            start, end = map(int, part.split('-'))
                            target_pages.update(range(start, end + 1))
                        except ValueError: pass
                    else:
                        try: target_pages.add(int(part))
                        except ValueError: pass
            else:
                target_pages = set(range(1, total_pages + 1))

            filtered_pages = set()
            for p in target_pages:
                if split_parity == "even" and p % 2 != 0: continue
                if split_parity == "odd" and p % 2 == 0: continue
                filtered_pages.add(p)

            pages_to_keep = set(range(1, total_pages + 1)) - filtered_pages if split_mode == "remove" else filtered_pages

            for p in sorted(list(pages_to_keep)):
                if 1 <= p <= total_pages:
                    writer.add_page(reader.pages[p - 1])

            if len(writer.pages) == 0: raise ValueError("No pages left to keep after applying filters.")

            output_path = os.path.join(user_output_dir(request.user), f"split_{uuid.uuid4().hex}.pdf")
            with open(output_path, "wb") as out:
                writer.write(out)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + "_split.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="Split PDF",
                description=f"Split {uploaded.name} keeping/removing specific pages",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "Split PDF", "split_mode": split_mode, "split_pages": split_pages_str, "split_parity": split_parity, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "PDF split successfully.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"Split failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            cleanup_file(temp_pdf)


# ═════════════════════════════════════════════════════════════════
# PDF COLOR ENHANCE
# ═════════════════════════════════════════════════════════════════


class PDFColorEnhanceView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    ALLOWED_EFFECTS = {"whiten", "grayscale", "high_contrast", "sharpen", "denoise"}

    @staticmethod
    def adaptive_black_white(gray, intensity):
        gray = gray.convert("L")
        histogram = gray.histogram()
        total_pixels = max(sum(histogram), 1)
        cumulative = 0
        median_value = 128
        for value, count in enumerate(histogram):
            cumulative += count
            if cumulative >= total_pixels // 2:
                median_value = value
                break
        if median_value < 110: gray = ImageOps.invert(gray)
        gray = ImageOps.autocontrast(gray, cutoff=(1, 1))
        contrast_factor = 2.4 + ((intensity - 10) / 90.0) * 3.2
        gray = ImageEnhance.Contrast(gray).enhance(contrast_factor)
        threshold = max(175, min(200, int(175 + ((intensity - 10) / 90.0) * 25)))
        bw = gray.point(lambda pixel: 0 if pixel < threshold else 255, mode="1")
        return bw.convert("RGB")

    @staticmethod
    def full_grayscale(image, intensity):
        gray = ImageOps.grayscale(image.convert("RGB"))
        histogram = gray.histogram()
        total_pixels = max(sum(histogram), 1)
        cumulative = 0
        median_value = 128
        for value, count in enumerate(histogram):
            cumulative += count
            if cumulative >= total_pixels // 2:
                median_value = value
                break
        if median_value < 110:
            gray = gray.point(lambda p: int(45 + (p / 255.0) * 190))
        else:
            gray = ImageOps.autocontrast(gray, cutoff=(1, 1))
        contrast_factor = 1.0 + ((intensity - 10) / 90.0) * 0.45
        gray = ImageEnhance.Contrast(gray).enhance(contrast_factor)
        return Image.merge("RGB", (gray, gray, gray))

    @staticmethod
    def strong_high_contrast(gray, intensity):
        gray = gray.convert("L")
        histogram = gray.histogram()
        total_pixels = max(sum(histogram), 1)
        cumulative = 0
        median_value = 128
        for value, count in enumerate(histogram):
            cumulative += count
            if cumulative >= total_pixels // 2:
                median_value = value
                break
        if median_value < 110: gray = ImageOps.invert(gray)
        gray = ImageOps.autocontrast(gray, cutoff=(1, 1))
        contrast_factor = 3.0 + ((intensity - 10) / 90.0) * 4.5
        gray = ImageEnhance.Contrast(gray).enhance(contrast_factor)
        brightness_factor = 1.10 + ((intensity - 10) / 90.0) * 0.22
        gray = ImageEnhance.Brightness(gray).enhance(brightness_factor)
        threshold = max(182, min(195, int(182 + ((intensity - 10) / 90.0) * 13)))
        binary = gray.point(lambda pixel: 0 if pixel < threshold else 255, mode="1")
        return binary.convert("RGB")

    @staticmethod
    def sharpen_text(image, intensity):
        image = image.convert("RGB")
        sharpness_factor = 1.7 + ((intensity - 10) / 90.0) * 3.8
        image = ImageEnhance.Sharpness(image).enhance(sharpness_factor)
        image = image.filter(ImageFilter.UnsharpMask(radius=1.0 + intensity / 180.0, percent=int(160 + intensity * 2), threshold=max(1, int(5 - intensity / 30))))
        return image

    @staticmethod
    def denoise_scan(image, intensity):
        image = image.convert("RGB")
        if intensity < 45: image = image.filter(ImageFilter.MedianFilter(size=3))
        elif intensity < 80: image = image.filter(ImageFilter.MedianFilter(size=5))
        else: image = image.filter(ImageFilter.MedianFilter(size=5)).filter(ImageFilter.GaussianBlur(radius=0.20))
        return image

    def post(self, request):
        if "file" not in request.FILES:
            return Response({"success": False, "message": "PDF file required."}, status=status.HTTP_400_BAD_REQUEST)

        uploaded = request.FILES["file"]
        if not uploaded.name.lower().endswith(".pdf"):
            return Response({"success": False, "message": "Only PDF files are supported."}, status=status.HTTP_400_BAD_REQUEST)

        effects_raw = request.data.get("effects") or request.data.get("color_mode") or ""
        raw_effects = effects_raw if isinstance(effects_raw, list) else str(effects_raw).split(",")
        effects = [str(effect).strip().lower() for effect in raw_effects if str(effect).strip().lower() in self.ALLOWED_EFFECTS]

        intensity = clamp_int(request.data.get("intensity", 50), 10, 100, 50)
        if not effects:
            return Response({"success": False, "message": "Please select at least one enhancement option."}, status=status.HTTP_400_BAD_REQUEST)

        operation = CombineOperation.objects.create(user=request.user, operation_type="pdf_color_enhance", input_count=1, options={"effects": effects, "intensity": intensity})
        temp_pdf = None
        source_doc = None
        output_doc = None
        temp_pages = []

        try:
            temp_pdf = save_uploaded_temp(uploaded)
            source_doc = fitz.open(temp_pdf)
            output_doc = fitz.open()
            if source_doc.page_count == 0: raise ValueError("The PDF contains no pages.")

            render_dpi = 200
            for page in source_doc:
                grayscale_selected = "grayscale" in effects
                if grayscale_selected:
                    pix = page.get_pixmap(dpi=render_dpi, colorspace=fitz.csGRAY, alpha=False, annots=True)
                    image = Image.frombytes("L", (pix.width, pix.height), pix.samples)
                else:
                    pix = page.get_pixmap(dpi=render_dpi, colorspace=fitz.csRGB, alpha=False, annots=True)
                    image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

                if "denoise" in effects: image = self.denoise_scan(image, intensity)
                if "grayscale" in effects: image = self.full_grayscale(image, intensity)
                if "whiten" in effects: image = self.adaptive_black_white(ImageOps.grayscale(image), intensity)
                if "high_contrast" in effects: image = self.strong_high_contrast(ImageOps.grayscale(image), intensity)
                if "sharpen" in effects: image = self.sharpen_text(image, intensity)

                image = image.convert("RGB")
                page_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_enhanced.jpg")
                temp_pages.append(page_path)
                image.save(page_path, "JPEG", quality=96, optimize=True)

                output_page = output_doc.new_page(width=page.rect.width, height=page.rect.height)
                output_page.insert_image(output_page.rect, filename=page_path, keep_proportion=False)
                image.close()
                cleanup_file(page_path)

            output_path = os.path.join(user_output_dir(request.user), f"enhanced_{uuid.uuid4().hex}.pdf")
            output_doc.save(output_path, garbage=4, deflate=True, clean=True)
            output_doc.close()
            source_doc.close()
            cleanup_file(temp_pdf)

            operation.output_file = relative_media_path(output_path)
            operation.output_name = os.path.splitext(uploaded.name)[0] + "_enhanced.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="PDF Color Enhance",
                description=f"Enhanced colors of {uploaded.name}",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "PDF Color Enhance", "effects": effects, "intensity": intensity, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": "PDF enhanced successfully.", "effects": effects, "intensity": intensity, "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id), "output_name": operation.output_name, "output_size": operation.output_size}, status=status.HTTP_200_OK)

        except Exception as exc:
            if output_doc is not None:
                try: output_doc.close()
                except Exception: pass
            if source_doc is not None:
                try: source_doc.close()
                except Exception: pass
            cleanup_file(temp_pdf)
            for path in temp_pages: cleanup_file(path)
            
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": "PDF enhancement failed.", "error": str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ═════════════════════════════════════════════════════════════════
# N-UP PDF
# ═════════════════════════════════════════════════════════════════


class NupPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        files = request.FILES.getlist("files")
        if not files and "file" in request.FILES:
            files = [request.FILES["file"]]
        if not files:
            return Response({"success": False, "message": "At least 1 PDF or image required."}, status=status.HTTP_400_BAD_REQUEST)

        pages_per_sheet = clamp_int(request.data.get("pages_per_sheet", 4), 2, 9, 4)
        spacing = request.data.get("spacing", "small")
        border = str(request.data.get("border", "false")).lower() == "true"
        layout = request.data.get("layout", "grid")
        page_size = request.data.get("page_size", "a4")
        orientation = request.data.get("orientation", "portrait")
        margin = request.data.get("margin", "medium")
        fit_mode = request.data.get("fit_mode", "fit")
        border_color = request.data.get("border_color", "#64748b")
        bg_color = request.data.get("bg_color", "#ffffff")

        operation = CombineOperation.objects.create(
            user=request.user, operation_type="nup_pdf", input_count=len(files),
            options={"pages_per_sheet": pages_per_sheet, "spacing": spacing, "border": border, "layout": layout, "page_size": page_size, "orientation": orientation, "margin": margin, "fit_mode": fit_mode},
        )

        try:
            source_images = []
            for uploaded in files:
                data = uploaded.read()
                if uploaded.name.lower().endswith(".pdf"):
                    d = fitz.open(stream=data, filetype="pdf")
                    try:
                        for p in d:
                            pix = p.get_pixmap(dpi=150)
                            source_images.append((pix.width, pix.height, pix.tobytes("png")))
                    finally:
                        d.close()
                else:
                    with Image.open(BytesIO(data)) as source:
                        img = source.convert("RGB")
                        buffer = BytesIO()
                        img.save(buffer, format="PNG")
                        source_images.append((img.width, img.height, buffer.getvalue()))

            if not source_images: raise ValueError("No valid pages/images found.")

            sizes = {"a4": (595, 842), "letter": (612, 792), "a3": (842, 1191)}
            pw, ph = sizes.get(page_size, sizes["a4"])
            if orientation == "landscape": pw, ph = ph, pw

            m = {"none": 0, "small": 12, "medium": 24, "large": 40}.get(margin, 24)
            gap = {"none": 0, "small": 4, "medium": 8, "large": 16}.get(spacing, 4)

            if layout == "horizontal": cols, rows = pages_per_sheet, 1
            elif layout == "vertical": cols, rows = 1, pages_per_sheet
            else: cols, rows = {2: (2, 1), 3: (3, 1), 4: (2, 2), 6: (3, 2), 9: (3, 3)}.get(pages_per_sheet, (2, 2))

            bg = hex_to_rgb01(bg_color)
            bc = hex_to_rgb01(border_color, (0.39, 0.43, 0.55))
            output_doc = fitz.open()

            for chunk_start in range(0, len(source_images), pages_per_sheet):
                chunk = source_images[chunk_start:chunk_start + pages_per_sheet]
                page = output_doc.new_page(width=pw, height=ph)
                page.draw_rect(page.rect, color=None, fill=bg)

                cell_w = (pw - 2 * m - (cols - 1) * gap) / cols
                cell_h = (ph - 2 * m - (rows - 1) * gap) / rows

                for index, (iw, ih, png) in enumerate(chunk):
                    if layout == "horizontal": row, col = 0, index
                    elif layout == "vertical": row, col = index, 0
                    else: row, col = index // cols, index % cols

                    x = m + col * (cell_w + gap)
                    y = m + row * (cell_h + gap)

                    if fit_mode == "stretch": rx, ry, rw, rh = x, y, cell_w, cell_h
                    else:
                        scale = min(cell_w / iw, cell_h / ih) if fit_mode == "fit" else max(cell_w / iw, cell_h / ih)
                        rw, rh = iw * scale, ih * scale
                        rx, ry = x + (cell_w - rw) / 2, y + (cell_h - rh) / 2

                    page.insert_image(fitz.Rect(rx, ry, rx + rw, ry + rh), stream=png)
                    if border:
                        page.draw_rect(fitz.Rect(x, y, x + cell_w, y + cell_h), color=bc, width=0.8)

            output_path = os.path.join(user_output_dir(request.user), f"nup_{pages_per_sheet}up_{uuid.uuid4().hex}.pdf")
            output_doc.save(output_path)
            output_doc.close()

            operation.output_file = relative_media_path(output_path)
            operation.output_name = f"nup_{pages_per_sheet}up.pdf"
            operation.output_size = os.path.getsize(output_path)
            operation.status = "completed"
            operation.completed_at = timezone.now()
            operation.save()

            # ✅ ADD UNIFIED HISTORY
            save_to_history(
                user=request.user,
                history_type="combine",
                title="N-up PDF",
                description=f"Combined {len(files)} files into {pages_per_sheet} pages per sheet",
                source_app="combine",
                source_model="CombineOperation",
                source_id=str(operation.id),
                metadata={"conversionType": "N-up PDF", "pages_per_sheet": pages_per_sheet, "layout": layout, "output": operation.output_name, "size": operation.output_size},
                output_file=operation.output_file,
                status="completed"
            )

            return Response({"success": True, "message": f"{pages_per_sheet} pages/images merged per sheet.", "download_url": f"/api/combine/{operation.id}/download/", "operation_id": str(operation.id)})
        except Exception as exc:
            operation.status = "failed"
            operation.error_message = str(exc)
            operation.save(update_fields=["status", "error_message"])
            return Response({"success": False, "message": f"N-up merge failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ═════════════════════════════════════════════════════════════════
# HISTORY
# ═════════════════════════════════════════════════════════════════


class CombineListView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        operation_type = request.query_params.get("type")
        queryset = CombineOperation.objects.filter(user=request.user).order_by("-created_at")
        if operation_type:
            queryset = queryset.filter(operation_type=operation_type)

        page = max(1, clamp_int(request.query_params.get("page", 1), 1, 10_000_000, 1))
        page_size = clamp_int(request.query_params.get("page_size", 10), 1, 100, 10)
        start = (page - 1) * page_size
        end = page * page_size
        total = queryset.count()
        serializer = CombineOperationSerializer(queryset[start:end], many=True)

        return Response({"success": True, "total": total, "page": page, "page_size": page_size, "data": serializer.data})


# ═════════════════════════════════════════════════════════════════
# DOWNLOAD
# ═════════════════════════════════════════════════════════════════


@api_view(["GET"])
@permission_classes([permissions.IsAuthenticated])
def download_combined(request, pk):
    try:
        operation = CombineOperation.objects.get(pk=pk, user=request.user)
        if not operation.output_file:
            return Response({"success": False, "message": "File not ready."}, status=status.HTTP_400_BAD_REQUEST)

        file_path = operation.output_file.path
        if not os.path.exists(file_path):
            return Response({"success": False, "message": "File not found."}, status=status.HTTP_404_NOT_FOUND)

        return FileResponse(open(file_path, "rb"), as_attachment=True, filename=operation.output_name or os.path.basename(file_path))
    except CombineOperation.DoesNotExist:
        return Response({"success": False, "message": "Operation not found."}, status=status.HTTP_404_NOT_FOUND)


# ═════════════════════════════════════════════════════════════════
# DELETE
# ═════════════════════════════════════════════════════════════════


@api_view(["DELETE"])
@permission_classes([permissions.IsAuthenticated])
def delete_combined(request, pk):
    try:
        operation = CombineOperation.objects.get(pk=pk, user=request.user)
        if operation.output_file:
            try:
                file_path = operation.output_file.path
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)
            except (ValueError, AttributeError, OSError):
                pass
        operation.delete()
        return Response({"success": True, "message": "Deleted successfully."})
    except CombineOperation.DoesNotExist:
        return Response({"success": False, "message": "Operation not found."}, status=status.HTTP_404_NOT_FOUND)