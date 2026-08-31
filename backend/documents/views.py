import os
import io
import uuid
import re
import logging

from django.conf import settings
from django.http import FileResponse, Http404

from rest_framework import status, permissions
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.views import APIView

from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

try:
    from reportlab.lib.utils import simpleSplit
except Exception:
    simpleSplit = None

from html import escape

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

from .models import DocumentUpload, DocumentTemplate
from .serializers import DocumentUploadSerializer, DocumentTemplateSerializer
from translation.views import TranslationService
from history.utils import save_to_history


# ============================================================
# LOGGER
# ============================================================

logger = logging.getLogger(__name__)


# ============================================================
# HELPER: TEXT CLEANING
# ============================================================

def clean_extracted_text(text):
    """
    Clean document text while keeping it suitable for normal
    single-document processing.
    """
    if not text:
        return ""

    text = re.sub(
        r"[\x00-\x08\x0b\x0e-\x1f\x7f-\x9f]",
        "",
        str(text),
    )

    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines = [
        re.sub(r"[ \t]+", " ", line).strip()
        for line in text.split("\n")
    ]

    lines = [line for line in lines if line]

    return " ".join(lines).strip()


def clean_page_text(text):
    """
    Clean a single PDF/document page without destroying its
    page boundary or all internal line breaks.
    """
    if not text:
        return ""

    text = re.sub(
        r"[\x00-\x08\x0b\x0e-\x1f\x7f-\x9f]",
        "",
        str(text),
    )

    text = text.replace("\r\n", "\n").replace("\r", "\n")

    cleaned_lines = []
    for line in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", line).strip()
        if line:
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()


# ============================================================
# TEXT / PAGE EXTRACTION
# ============================================================

class DocumentProcessor:

    @staticmethod
    def extract_pages(file_path, file_type):
        """
        Return one item per real source page wherever the file format
        has a real page concept.
        """
        if file_type == "pdf":
            if fitz is not None:
                pdf = fitz.open(file_path)
                try:
                    return [
                        clean_page_text(page.get_text("text") or "")
                        for page in pdf
                    ]
                finally:
                    pdf.close()

            reader = PdfReader(file_path)
            pages = []
            for page_number, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception as exc:
                    logger.warning(
                        "PDF text extraction failed on page %s: %s",
                        page_number,
                        exc,
                    )
                    page_text = ""
                pages.append(clean_page_text(page_text))
            return pages

        if file_type in {"txt", "md", "rtf"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                text = file.read()
            parts = re.split(r"---\s*PAGE\s*BREAK\s*---|\f", text, flags=re.IGNORECASE)
            return [clean_page_text(part) for part in (parts or [text])]

        if file_type == "docx":
            document = DocxDocument(file_path)
            page_groups = [[]]
            explicit_break_found = False

            for paragraph in document.paragraphs:
                paragraph_text = paragraph.text or ""
                if paragraph_text.strip():
                    page_groups[-1].append(paragraph_text.strip())

                paragraph_has_break = False
                for run in paragraph.runs:
                    xml = run._element.xml
                    if 'w:type="page"' in xml or "w:type='page'" in xml:
                        paragraph_has_break = True
                        explicit_break_found = True
                        break

                if paragraph_has_break:
                    page_groups.append([])

            if explicit_break_found:
                return [clean_page_text("\n".join(group)) for group in page_groups]

            text = "\n".join(
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            )
            return [clean_page_text(text)]

        if file_type == "html":
            from html.parser import HTMLParser

            class MLStripper(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.fed = []

                def handle_data(self, data):
                    if data and data.strip():
                        self.fed.append(data.strip())

                def get_data(self):
                    return "\n".join(self.fed)

            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                parser = MLStripper()
                parser.feed(file.read())
                text = parser.get_data()
            return [clean_page_text(text)]

        raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def extract_text(file_path, file_type):
        pages = DocumentProcessor.extract_pages(file_path, file_type)
        return "\n\n".join(page for page in pages if page).strip()

    # ========================================================
    # PDF FONT HELPERS
    # ========================================================

    @staticmethod
    def _find_unicode_font_file():
        candidates = [
            r"C:\Windows\Fonts\Nirmala.ttf",
            r"C:\Windows\Fonts\NirmalaUI.ttf",
            r"C:\Windows\Fonts\segoeui.ttf",
            r"C:\Windows\Fonts\arial.ttf",
            r"/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
            r"/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            r"/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
            r"/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]
        for path in candidates:
            if os.path.exists(path):
                return path
        return None

    @staticmethod
    def _fit_font_size_for_rect(page, text, rect, original_size, fontfile=None):
        """
        Preserve the source font size whenever possible.
        Only reduce it when the translated text cannot fit inside the
        original text rectangle.
        """
        size = max(float(original_size or 10), 4.0)
        min_size = max(size * 0.72, 6.0)
        test_rect = fitz.Rect(rect)
        test_rect.y1 += max(2.0, size * 0.45)

        while size >= min_size:
            try:
                rc = page.insert_textbox(
                    test_rect,
                    text,
                    fontsize=size,
                    fontfile=fontfile,
                    fontname="QuillFont" if fontfile else "helv",
                    color=(0, 0, 0),
                    overlay=True,
                    render_mode=0,
                )
                # If return value is >= 0, the text fitted.
                # Remove the just-created content by redacting later; this is only
                # a measurement pass, so we use a scratch page approach below.
                return size
            except Exception:
                pass
            size -= 0.5

        return min_size

    @staticmethod
    def _translate_pdf_preserve_layout(
        original_pdf_path,
        translated_pages,
        output_path,
    ):
        """
        Translate a PDF while preserving the REAL page count, source
        page dimensions, approximate text positions and ORIGINAL FONT SIZES.

        For text-based PDFs, each text line is translated independently and
        reinserted into its original bounding box using the source font size.
        Font size is reduced only when the translated text cannot fit.

        Image-only/scanned pages are kept unchanged because there is no
        reliable text layer to translate without OCR.
        """
        if fitz is None:
            raise RuntimeError(
                "PyMuPDF is required for layout-preserving PDF translation. "
                "Install it with: pip install pymupdf"
            )

        source_pdf = fitz.open(original_pdf_path)
        output_pdf = fitz.open()
        fontfile = DocumentProcessor._find_unicode_font_file()

        try:
            if len(source_pdf) != len(translated_pages):
                raise RuntimeError(
                    "Source page count and translated page count do not match."
                )

            for page_index, source_page in enumerate(source_pdf):
                translated_text = translated_pages[page_index] or ""
                output_page = output_pdf.new_page(
                    width=source_page.rect.width,
                    height=source_page.rect.height,
                )

                # Copy original page as vector PDF content.
                # This keeps images, backgrounds, borders and page dimensions.
                output_page.show_pdf_page(
                    output_page.rect,
                    source_pdf,
                    page_index,
                )

                blocks = source_page.get_text("dict").get("blocks", [])
                source_lines = []

                for block in blocks:
                    if block.get("type") != 0:
                        continue

                    for line in block.get("lines", []):
                        spans = [
                            span
                            for span in line.get("spans", [])
                            if str(span.get("text", "")).strip()
                        ]
                        if not spans:
                            continue

                        line_text = "".join(
                            str(span.get("text", ""))
                            for span in spans
                        ).strip()
                        if not line_text:
                            continue

                        rect = fitz.Rect(line.get("bbox", block.get("bbox")))
                        size_values = [
                            float(span.get("size", 10) or 10)
                            for span in spans
                        ]
                        original_size = max(size_values)

                        source_lines.append({
                            "text": line_text,
                            "rect": rect,
                            "size": original_size,
                        })

                # If no text layer exists, preserve this page as-is.
                if not source_lines:
                    logger.warning(
                        "Page %s has no selectable text. Preserving original page.",
                        page_index + 1,
                    )
                    continue

                # We need to map translated content back to the same text lines.
                # The page translation endpoint already returns page-level text.
                # Use the translated lines in order. If translation changed the
                # number of lines, distribute by source line count rather than
                # creating additional pages.
                translated_lines = [
                    line.strip()
                    for line in str(translated_text).splitlines()
                    if line.strip()
                ]

                if not translated_lines:
                    # Fallback to original page content.
                    continue

                # Create a stable mapping. Prefer one-to-one when counts match.
                mapped = []
                if len(translated_lines) == len(source_lines):
                    mapped = translated_lines
                elif len(source_lines) == 1:
                    mapped = [" ".join(translated_lines)]
                else:
                    # Use proportional grouping so we do not lose translated text.
                    total = len(translated_lines)
                    count = len(source_lines)
                    cursor = 0
                    for i in range(count):
                        remaining_lines = total - cursor
                        remaining_blocks = count - i
                        take = max(1, round(remaining_lines / remaining_blocks))
                        if i == count - 1:
                            take = remaining_lines
                        piece = " ".join(translated_lines[cursor:cursor + take]).strip()
                        mapped.append(piece)
                        cursor += take

                # Redact only original text rectangles. The underlying page layout,
                # images and borders remain untouched outside these small regions.
                for item in source_lines:
                    rect = fitz.Rect(item["rect"])
                    # Slight expansion gives translated glyphs enough vertical room.
                    rect.x0 -= 1.0
                    rect.x1 += 1.0
                    rect.y0 -= 1.0
                    rect.y1 += max(2.0, item["size"] * 0.35)
                    output_page.add_redact_annot(
                        rect,
                        fill=(1, 1, 1),
                    )

                output_page.apply_redactions(
                    images=fitz.PDF_REDACT_IMAGE_NONE,
                    graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                    text=fitz.PDF_REDACT_TEXT_REMOVE,
                )

                # Insert translated lines using the ORIGINAL FONT SIZE.
                for idx, item in enumerate(source_lines):
                    if idx >= len(mapped):
                        break

                    text = mapped[idx]
                    rect = fitz.Rect(item["rect"])
                    rect.x1 += max(6.0, rect.width * 0.08)
                    rect.y1 += max(3.0, item["size"] * 0.5)

                    original_size = float(item["size"] or 10)
                    chosen_size = original_size
                    min_size = max(original_size * 0.72, 6.0)

                    # Try original size first. Then reduce only if needed.
                    while chosen_size >= min_size:
                        try:
                            result = output_page.insert_textbox(
                                rect,
                                text,
                                fontsize=chosen_size,
                                fontfile=fontfile,
                                fontname="QuillFont" if fontfile else "helv",
                                color=(0, 0, 0),
                                align=fitz.TEXT_ALIGN_LEFT,
                                overlay=True,
                            )
                            if result >= 0:
                                break
                        except Exception as exc:
                            logger.debug(
                                "Insert text retry on page %s line %s: %s",
                                page_index + 1,
                                idx + 1,
                                exc,
                            )
                        chosen_size -= 0.5

                    if chosen_size < min_size:
                        try:
                            output_page.insert_textbox(
                                rect,
                                text,
                                fontsize=min_size,
                                fontfile=fontfile,
                                fontname="QuillFont" if fontfile else "helv",
                                color=(0, 0, 0),
                                align=fitz.TEXT_ALIGN_LEFT,
                                overlay=True,
                            )
                        except Exception as exc:
                            logger.warning(
                                "Could not insert translated text on page %s line %s: %s",
                                page_index + 1,
                                idx + 1,
                                exc,
                            )

            output_pdf.save(
                output_path,
                garbage=4,
                deflate=True,
                clean=True,
            )

        finally:
            output_pdf.close()
            source_pdf.close()

        return output_path

    # ========================================================
    # EXACT-PAGE FALLBACK PDF OUTPUT
    # ========================================================

    @staticmethod
    def _register_pdf_font():
        candidates = [
            ("QuillUnicode", r"C:\Windows\Fonts\Nirmala.ttf"),
            ("QuillUnicode", r"C:\Windows\Fonts\NirmalaUI.ttf"),
            ("QuillUnicode", r"C:\Windows\Fonts\segoeui.ttf"),
            ("QuillUnicode", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
            ("QuillUnicode", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            ("QuillUnicode", "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
        ]
        for name, path in candidates:
            if os.path.exists(path):
                try:
                    if name not in pdfmetrics.getRegisteredFontNames():
                        pdfmetrics.registerFont(TTFont(name, path))
                    return name
                except Exception:
                    continue
        return "Helvetica"

    @staticmethod
    def _wrap_page_text(text, font_name, font_size, max_width):
        paragraphs = str(text or "").splitlines() or [""]
        lines = []
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                lines.append("")
                continue
            words = paragraph.split()
            current = ""
            for word in words:
                candidate = word if not current else f"{current} {word}"
                if pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
                    current = candidate
                else:
                    if current:
                        lines.append(current)
                    current = word
            if current:
                lines.append(current)
        return lines

    @staticmethod
    def _draw_fitted_page(pdf, text, page_width, page_height):
        margin_x = 40
        margin_y = 40
        font_name = DocumentProcessor._register_pdf_font()
        clean_text = str(text or "").strip()
        size = 12
        while size >= 4:
            lines = DocumentProcessor._wrap_page_text(
                clean_text,
                font_name,
                size,
                max(page_width - margin_x * 2, 20),
            )
            if len(lines) * max(size * 1.25, 6) <= max(page_height - margin_y * 2, 20):
                break
            size -= 0.5
        pdf.setFont(font_name, size)
        y = page_height - margin_y
        leading = max(size * 1.25, 5)
        for line in lines:
            if y < margin_y:
                break
            if line:
                pdf.drawString(margin_x, y, line)
            y -= leading

    @staticmethod
    def create_pdf_from_pages(translated_pages, output_path, original_pdf_path=None):
        """
        Exact page-count fallback. When source PDF is available and PyMuPDF is
        installed, use layout-preserving generation with original font sizes.
        """
        if original_pdf_path and fitz is not None and os.path.exists(original_pdf_path):
            return DocumentProcessor._translate_pdf_preserve_layout(
                original_pdf_path,
                translated_pages,
                output_path,
            )

        # Hard fallback if PyMuPDF is unavailable.
        page_sizes = []
        if original_pdf_path and os.path.exists(original_pdf_path):
            try:
                reader = PdfReader(original_pdf_path)
                for page in reader.pages:
                    page_sizes.append((float(page.mediabox.width), float(page.mediabox.height)))
            except Exception as exc:
                logger.warning("Could not read original PDF page sizes: %s", exc)

        pdf = canvas.Canvas(output_path, pagesize=letter)
        for index, page_text in enumerate(translated_pages):
            width, height = page_sizes[index] if index < len(page_sizes) else letter
            pdf.setPageSize((width, height))
            DocumentProcessor._draw_fitted_page(pdf, page_text, width, height)
            pdf.showPage()
        pdf.save()
        return output_path

    # ========================================================
    # GENERAL OUTPUT
    # ========================================================

    @staticmethod
    def create_output(text, output_format, output_path, original_name, pages=None, original_pdf_path=None):
        output_format = str(output_format or "pdf").lower().strip()

        if output_format == "txt":
            final_text = (
                "\n\n".join(
                    f"--- PAGE {i + 1} ---\n{page}"
                    for i, page in enumerate(pages)
                )
                if pages
                else clean_extracted_text(text)
            )
            with open(output_path, "w", encoding="utf-8") as file:
                file.write(final_text)
            return output_path

        if output_format == "docx":
            document = DocxDocument()
            if pages:
                for index, page_text in enumerate(pages):
                    if page_text:
                        document.add_paragraph(page_text)
                    if index < len(pages) - 1:
                        paragraph = document.add_paragraph()
                        paragraph.add_run().add_break(WD_BREAK.PAGE)
            else:
                document.add_paragraph(clean_extracted_text(text))
            document.save(output_path)
            return output_path

        if output_format == "pdf":
            return DocumentProcessor.create_pdf_from_pages(
                pages if pages else [clean_extracted_text(text)],
                output_path,
                original_pdf_path=original_pdf_path,
            )

        if output_format == "html":
            if pages:
                sections = []
                for index, page_text in enumerate(pages):
                    page_html = escape(page_text or "").replace("\n", "<br>")
                    sections.append(
                        f"""
                        <section class="doc-page">
                            <div class="page-number">Page {index + 1}</div>
                            <div class="page-content">{page_html or "&nbsp;"}</div>
                        </section>
                        """
                    )
                body = "".join(sections)
            else:
                body = (
                    f'<section class="doc-page"><div class="page-content">'
                    f'{escape(clean_extracted_text(text))}</div></section>'
                )

            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{escape(original_name)}</title>
<style>
body {{ margin:0; background:#e5e7eb; font-family:Arial,sans-serif; color:#1e293b; }}
.doc-page {{ width:210mm; min-height:297mm; box-sizing:border-box; background:#fff; margin:20px auto; padding:40px; position:relative; page-break-after:always; break-after:page; overflow:hidden; }}
.doc-page:last-child {{ page-break-after:auto; break-after:auto; }}
.page-number {{ font-size:10px; color:#94a3b8; margin-bottom:18px; }}
.page-content {{ font-size:12px; line-height:1.6; white-space:normal; overflow-wrap:anywhere; }}
@media print {{ body {{ background:#fff; }} .doc-page {{ margin:0; width:auto; min-height:0; height:297mm; }} }}
</style>
</head>
<body>{body}</body>
</html>"""
            with open(output_path, "w", encoding="utf-8") as file:
                file.write(html_content)
            return output_path

        raise ValueError(f"Unsupported output format: {output_format}")


# ============================================================
# DOCUMENT UPLOAD + TRANSLATION
# ============================================================

class DocumentUploadView(APIView):

    permission_classes = [
        permissions.IsAuthenticated
    ]

    def post(self, request):

        if "file" not in request.FILES:
            return Response(
                {
                    "success": False,
                    "message": "No file provided.",
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        uploaded_file = request.FILES[
            "file"
        ]

        file_type = (
            self._detect_file_type(
                uploaded_file.name
            )
        )

        if not file_type:
            return Response(
                {
                    "success": False,
                    "message": "Unsupported file type.",
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        max_file_size = getattr(
            settings,
            "MAX_DOCUMENT_UPLOAD_SIZE",
            250 * 1024 * 1024,
        )

        if uploaded_file.size > max_file_size:
            return Response(
                {
                    "success": False,
                    "message": "File too large.",
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        user = request.user

        source_lang = (
            request.data.get(
                "source_lang",
                "auto",
            )
            or "auto"
        )

        target_lang = (
            request.data.get(
                "target_lang",
                "en",
            )
            or "en"
        )

        output_format = (
            request.data.get(
                "output_format",
                "pdf",
            )
            or "pdf"
        ).lower()

        if output_format not in {
            "txt",
            "docx",
            "pdf",
            "html",
        }:
            output_format = "pdf"

        doc = DocumentUpload.objects.create(
            user=user,
            original_file=uploaded_file,
            original_name=uploaded_file.name,
            file_type=file_type,
            file_size=uploaded_file.size,
            source_lang=source_lang,
            target_lang=target_lang,
            output_format=output_format,
            preserve_formatting=True,
        )

        try:
            self._process_document(
                doc
            )

        except Exception as exc:
            logger.exception(
                "Document processing error: %s",
                exc,
            )

            doc.mark_failed(
                str(exc)
            )

            return Response(
                {
                    "success": False,
                    "message": str(exc),
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        return Response(
            {
                "success": True,
                "message": (
                    "Document translated successfully."
                ),
                "data": DocumentUploadSerializer(
                    doc
                ).data,
            },
            status=status.HTTP_200_OK,
        )

    def _detect_file_type(
        self,
        filename,
    ):
        extension = (
            filename.rsplit(
                ".",
                1,
            )[-1].lower()
            if "." in filename
            else ""
        )

        return {
            "pdf": "pdf",
            "docx": "docx",
            "doc": "docx",
            "txt": "txt",
            "rtf": "rtf",
            "html": "html",
            "htm": "html",
            "md": "md",
        }.get(extension)

    # ========================================================
    # PROCESS DOCUMENT
    # ========================================================

    def _process_document(
        self,
        doc,
    ):
        doc.mark_processing()

        original_path = (
            doc.original_file.path
        )

        # ----------------------------------------------------
        # PDF: REAL PAGE-BY-PAGE PROCESSING
        # ----------------------------------------------------

        if doc.file_type == "pdf":

            pages = (
                DocumentProcessor.extract_pages(
                    original_path,
                    "pdf",
                )
            )

            # The REAL number of pages in the uploaded PDF.
            real_page_count = len(
                pages
            )

            if real_page_count < 1:
                raise ValueError(
                    "PDF contains no pages."
                )

            print(
                f"📄 Original PDF pages: "
                f"{real_page_count}"
            )

            # Store page-preserving extracted text.
            extracted_text = "\n\n".join(
                page for page in pages
            )

            doc.extracted_text = (
                extracted_text[:10000]
            )

            doc.page_count = (
                real_page_count
            )

            doc.save(
                update_fields=[
                    "extracted_text",
                    "page_count",
                ]
            )

            # Translate each original page separately.
            translated_pages = []

            for index, page_text in enumerate(
                pages,
                start=1,
            ):

                print(
                    f"🔄 Translating PDF page "
                    f"{index}/{real_page_count}..."
                )

                if not page_text.strip():
                    # Keep empty/scanned-only page.
                    translated_pages.append("")
                    continue

                translated_page = (
                    self._translate_long_text(
                        page_text,
                        doc.source_lang,
                        doc.target_lang,
                    )
                )

                if not translated_page:
                    # Never remove a source page.
                    translated_page = page_text

                translated_pages.append(
                    translated_page
                )

            translated_full_text = (
                "\n\n".join(
                    translated_pages
                )
            )

            if not translated_full_text.strip():
                # A PDF may contain image-only pages.
                # We still create an exact-page output instead
                # of turning a 5-page PDF into a 1-page file.
                logger.warning(
                    "PDF has no extractable text. "
                    "Preserving all %s pages.",
                    real_page_count,
                )

            doc.translated_text_preview = (
                translated_full_text[:2000]
            )

            doc.save(
                update_fields=[
                    "translated_text_preview"
                ]
            )

            safe_name = self._safe_name(
                doc.original_name
            )

            output_path = os.path.join(
                settings.MEDIA_ROOT,
                "documents",
                str(doc.user.id),
                (
                    f"{uuid.uuid4().hex}."
                    f"{doc.output_format}"
                ),
            )

            os.makedirs(
                os.path.dirname(output_path),
                exist_ok=True,
            )

            if doc.output_format == "pdf":
                # EXACT SAME NUMBER OF PAGES.
                DocumentProcessor.create_pdf_from_pages(
                    translated_pages,
                    output_path,
                    original_pdf_path=original_path,
                )

            else:
                # Also preserve page boundaries in non-PDF outputs.
                DocumentProcessor.create_output(
                    translated_full_text,
                    doc.output_format,
                    output_path,
                    doc.original_name,
                    pages=translated_pages,
                    original_pdf_path=original_path,
                )

            relative_path = os.path.relpath(
                output_path,
                settings.MEDIA_ROOT,
            ).replace("\\", "/")

            doc.mark_completed(
                relative_path,
                (
                    f"{safe_name}_translated."
                    f"{doc.output_format}"
                ),
            )

            self._save_history(
                doc,
                relative_path,
            )

            return

        # ----------------------------------------------------
        # NON-PDF DOCUMENTS
        # ----------------------------------------------------

        extracted = (
            DocumentProcessor.extract_text(
                original_path,
                doc.file_type,
            )
            or ""
        )

        print(
            f"📊 Extracted text length: "
            f"{len(extracted)} characters"
        )

        if not extracted.strip():
            raise ValueError(
                "No readable text found in the document."
            )

        # Non-PDF files do not necessarily have a physical
        # page concept. Keep the original behavior here.
        pages = (
            DocumentProcessor.extract_pages(
                original_path,
                doc.file_type,
            )
        )

        doc.extracted_text = (
            extracted[:10000]
        )

        doc.page_count = max(
            1,
            len(pages),
        )

        doc.save(
            update_fields=[
                "extracted_text",
                "page_count",
            ]
        )

        print(
            f"🔄 Starting translation: "
            f"{doc.source_lang} -> "
            f"{doc.target_lang}"
        )

        translated = (
            self._translate_long_text(
                extracted,
                doc.source_lang,
                doc.target_lang,
            )
        )

        if not translated or not translated.strip():
            raise ValueError(
                "Translation returned no text."
            )

        print(
            "✅ Translation successful. "
            f"Translated length: {len(translated)}"
        )

        doc.translated_text_preview = (
            translated[:2000]
        )

        doc.save(
            update_fields=[
                "translated_text_preview"
            ]
        )

        safe_name = self._safe_name(
            doc.original_name
        )

        output_path = os.path.join(
            settings.MEDIA_ROOT,
            "documents",
            str(doc.user.id),
            (
                f"{uuid.uuid4().hex}."
                f"{doc.output_format}"
            ),
        )

        os.makedirs(
            os.path.dirname(output_path),
            exist_ok=True,
        )

        DocumentProcessor.create_output(
            translated,
            doc.output_format,
            output_path,
            doc.original_name,
        )

        relative_path = os.path.relpath(
            output_path,
            settings.MEDIA_ROOT,
        ).replace("\\", "/")

        doc.mark_completed(
            relative_path,
            (
                f"{safe_name}_translated."
                f"{doc.output_format}"
            ),
        )

        self._save_history(
            doc,
            relative_path,
        )

    # ========================================================
    # SAFE NAME
    # ========================================================

    @staticmethod
    def _safe_name(
        filename,
    ):
        return (
            re.sub(
                r"[^a-zA-Z0-9_\- ]+",
                "_",
                os.path.splitext(filename)[0],
            ).strip()
            or "translated"
        )

    # ========================================================
    # HISTORY
    # ========================================================

    @staticmethod
    def _save_history(
        doc,
        relative_path,
    ):
        try:
            save_to_history(
                user=doc.user,
                history_type="documents",
                title="Document Translation",
                description=(
                    f"Translated "
                    f"{doc.original_name}"
                ),
                source_app="documents",
                source_model="DocumentUpload",
                source_id=str(doc.id),
                metadata={
                    "outputFormat": (
                        doc.output_format
                    ),
                    "pages": (
                        doc.page_count
                    ),
                    "size": doc.file_size,
                    "output": (
                        doc.translated_name
                    ),
                },
                output_file=relative_path,
                status="completed",
            )

        except Exception as exc:
            logger.warning(
                "History save error: %s",
                exc,
            )

    # ========================================================
    # LONG TEXT TRANSLATION
    # ========================================================

    def _translate_long_text(
        self,
        text,
        source_lang,
        target_lang,
    ):
        if not text or not text.strip():
            return ""

        src = (
            "auto"
            if source_lang == "auto"
            else source_lang
        )

        max_chunk = 3500

        # Keep paragraphs/line structure as much as possible.
        normalized = (
            str(text)
            .replace("\r\n", "\n")
            .replace("\r", "\n")
        )

        chunks = []
        current_chunk = ""

        for paragraph in normalized.split("\n"):
            paragraph = paragraph.strip()

            if not paragraph:
                if current_chunk:
                    chunks.append(
                        current_chunk
                    )
                    current_chunk = ""
                continue

            words = paragraph.split()

            for word in words:
                candidate = (
                    word
                    if not current_chunk
                    else f"{current_chunk} {word}"
                )

                if (
                    len(candidate)
                    <= max_chunk
                ):
                    current_chunk = candidate
                else:
                    if current_chunk:
                        chunks.append(
                            current_chunk
                        )

                    # Handle a single giant token.
                    if len(word) > max_chunk:
                        for start in range(
                            0,
                            len(word),
                            max_chunk,
                        ):
                            chunks.append(
                                word[
                                    start:
                                    start + max_chunk
                                ]
                            )
                        current_chunk = ""
                    else:
                        current_chunk = word

            if current_chunk:
                chunks.append(
                    current_chunk
                )
                current_chunk = ""

        if not chunks and normalized.strip():
            chunks.append(
                normalized.strip()
            )

        translated_chunks = []

        for i, chunk in enumerate(
            chunks,
            start=1,
        ):
            try:
                print(
                    f"Translating chunk "
                    f"{i}/{len(chunks)} "
                    f"(Original Length: "
                    f"{len(chunk)})..."
                )

                translated = (
                    TranslationService.translate(
                        text=chunk,
                        source_lang=src,
                        target_lang=target_lang,
                        mode="text_to_text",
                    )
                )

                translated_str = str(
                    translated or ""
                ).strip()

                print(
                    "DEBUG Translation API Response: "
                    f"{repr(translated_str[:100])}"
                )

                # Obvious API error detection.
                error_keywords = {
                    "LIMIT",
                    "ERROR",
                    "500",
                    "RATE",
                    "EXCEEDED",
                    "TOO MANY",
                    "<HTML>",
                    "<!DOCTYPE",
                }

                upper = (
                    translated_str.upper()
                )

                is_error = any(
                    keyword in upper
                    for keyword in error_keywords
                )

                min_expected_len = (
                    len(chunk) * 0.30
                )

                if (
                    translated_str
                    and not is_error
                    and (
                        len(translated_str)
                        >= min_expected_len
                    )
                ):
                    translated_chunks.append(
                        translated_str
                    )

                    print(
                        f"✅ Chunk {i} "
                        "translated successfully."
                    )

                else:
                    # NEVER delete the original content
                    # because a translation API returned bad data.
                    print(
                        f"⚠️ Chunk {i} "
                        "translation rejected. "
                        "Keeping original text."
                    )

                    translated_chunks.append(
                        chunk
                    )

            except Exception as exc:
                print(
                    f"❌ Chunk {i} "
                    f"translation failed: {exc}. "
                    "Keeping original."
                )

                translated_chunks.append(
                    chunk
                )

        final_text = " ".join(
            translated_chunks
        ).strip()

        print(
            "✅ Translation complete. "
            f"Final text length: {len(final_text)}"
        )

        return final_text


# ============================================================
# DOCUMENT LIST
# ============================================================

class DocumentListView(APIView):

    permission_classes = [
        permissions.IsAuthenticated
    ]

    def get(self, request):

        status_filter = (
            request.query_params.get(
                "status"
            )
        )

        queryset = (
            DocumentUpload.objects
            .filter(
                user=request.user
            )
            .order_by(
                "-created_at"
            )
        )

        if status_filter:
            queryset = queryset.filter(
                status=status_filter
            )

        try:
            page = max(
                int(
                    request.query_params.get(
                        "page",
                        1,
                    )
                ),
                1,
            )

            page_size = min(
                max(
                    int(
                        request.query_params.get(
                            "page_size",
                            10,
                        )
                    ),
                    1,
                ),
                100,
            )

        except (
            TypeError,
            ValueError,
        ):
            page = 1
            page_size = 10

        start = (
            page - 1
        ) * page_size

        serializer = (
            DocumentUploadSerializer(
                queryset[
                    start:
                    start + page_size
                ],
                many=True,
            )
        )

        return Response(
            {
                "success": True,
                "total": queryset.count(),
                "page": page,
                "page_size": page_size,
                "data": serializer.data,
            }
        )


# ============================================================
# DOCUMENT DETAIL
# ============================================================

class DocumentDetailView(APIView):

    permission_classes = [
        permissions.IsAuthenticated
    ]

    def get(
        self,
        request,
        pk,
    ):
        try:

            doc = (
                DocumentUpload.objects.get(
                    pk=pk,
                    user=request.user,
                )
            )

            return Response(
                {
                    "success": True,
                    "data": (
                        DocumentUploadSerializer(
                            doc
                        ).data
                    ),
                }
            )

        except DocumentUpload.DoesNotExist:

            return Response(
                {
                    "success": False,
                    "message": (
                        "Document not found."
                    ),
                },
                status=(
                    status.HTTP_404_NOT_FOUND
                ),
            )


# ============================================================
# DOWNLOAD TRANSLATED DOCUMENT
# ============================================================

@api_view(["GET"])
@permission_classes(
    [permissions.IsAuthenticated]
)
def download_document(
    request,
    pk,
):
    try:

        doc = (
            DocumentUpload.objects.get(
                pk=pk,
                user=request.user,
            )
        )

        if not doc.translated_file:
            return Response(
                {
                    "success": False,
                    "message": (
                        "Translation not yet complete."
                    ),
                },
                status=(
                    status.HTTP_400_BAD_REQUEST
                ),
            )

        if not os.path.exists(
            doc.translated_file.path
        ):
            raise Http404(
                "Translated file not found."
            )

        return FileResponse(
            open(
                doc.translated_file.path,
                "rb",
            ),
            as_attachment=True,
            filename=(
                doc.translated_name
                or (
                    f"translated."
                    f"{doc.output_format}"
                )
            ),
        )

    except DocumentUpload.DoesNotExist:

        return Response(
            {
                "success": False,
                "message": (
                    "Document not found."
                ),
            },
            status=(
                status.HTTP_404_NOT_FOUND
            ),
        )


# ============================================================
# DOWNLOAD ORIGINAL DOCUMENT
# ============================================================

@api_view(["GET"])
@permission_classes(
    [permissions.IsAuthenticated]
)
def download_original(
    request,
    pk,
):
    try:

        doc = (
            DocumentUpload.objects.get(
                pk=pk,
                user=request.user,
            )
        )

        if not os.path.exists(
            doc.original_file.path
        ):
            raise Http404(
                "Original file not found."
            )

        return FileResponse(
            open(
                doc.original_file.path,
                "rb",
            ),
            as_attachment=True,
            filename=doc.original_name,
        )

    except DocumentUpload.DoesNotExist:

        return Response(
            {
                "success": False,
                "message": (
                    "Document not found."
                ),
            },
            status=(
                status.HTTP_404_NOT_FOUND
            ),
        )


# ============================================================
# DELETE DOCUMENT
# ============================================================

@api_view(["DELETE"])
@permission_classes(
    [permissions.IsAuthenticated]
)
def delete_document(
    request,
    pk,
):
    try:

        doc = (
            DocumentUpload.objects.get(
                pk=pk,
                user=request.user,
            )
        )

        if (
            hasattr(
                doc.original_file,
                "path",
            )
            and os.path.exists(
                doc.original_file.path
            )
        ):
            os.remove(
                doc.original_file.path
            )

        if (
            hasattr(
                doc.translated_file,
                "path",
            )
            and os.path.exists(
                doc.translated_file.path
            )
        ):
            os.remove(
                doc.translated_file.path
            )

        doc.delete()

        return Response(
            {
                "success": True,
                "message": (
                    "Document deleted successfully."
                ),
            }
        )

    except DocumentUpload.DoesNotExist:

        return Response(
            {
                "success": False,
                "message": (
                    "Document not found."
                ),
            },
            status=(
                status.HTTP_404_NOT_FOUND
            ),
        )


# ============================================================
# DOCUMENT TEMPLATES
# ============================================================

class DocumentTemplateView(APIView):

    permission_classes = [
        permissions.IsAuthenticated
    ]

    def get(
        self,
        request,
    ):
        templates = (
            DocumentTemplate.objects
            .filter(
                user=request.user
            )
            .order_by(
                "-created_at"
            )
        )

        return Response(
            {
                "success": True,
                "data": (
                    DocumentTemplateSerializer(
                        templates,
                        many=True,
                    ).data
                ),
            }
        )

    def post(
        self,
        request,
    ):
        serializer = (
            DocumentTemplateSerializer(
                data=request.data
            )
        )

        if serializer.is_valid():

            serializer.save(
                user=request.user
            )

            return Response(
                {
                    "success": True,
                    "data": serializer.data,
                },
                status=(
                    status.HTTP_201_CREATED
                ),
            )

        return Response(
            {
                "success": False,
                "errors": (
                    serializer.errors
                ),
            },
            status=(
                status.HTTP_400_BAD_REQUEST
            ),
        )


# ============================================================
# DELETE TEMPLATE
# ============================================================

@api_view(["DELETE"])
@permission_classes(
    [permissions.IsAuthenticated]
)
def delete_template(
    request,
    pk,
):
    try:

        (
            DocumentTemplate.objects
            .get(
                pk=pk,
                user=request.user,
            )
            .delete()
        )

        return Response(
            {
                "success": True,
                "message": (
                    "Template deleted."
                ),
            }
        )

    except DocumentTemplate.DoesNotExist:

        return Response(
            {
                "success": False,
                "message": (
                    "Template not found."
                ),
            },
            status=(
                status.HTTP_404_NOT_FOUND
            ),
        )
